from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document(r'C:\UnityGames\GameYear\Diplom\Диплома.docx')

def style_table(table):
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = cell._tc.get_or_add_tcPr()
        shading_elem = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '2F5496',
            qn('w:val'): 'clear'
        })
        shading.append(shading_elem)
    for i, row in enumerate(table.rows[1:], 1):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(10)
            if i % 2 == 0:
                shading = cell._tc.get_or_add_tcPr()
                shading_elem = shading.makeelement(qn('w:shd'), {
                    qn('w:fill'): 'D6E4F0',
                    qn('w:val'): 'clear'
                })
                shading.append(shading_elem)

def find_paragraph_idx(contains_text):
    for i, p in enumerate(doc.paragraphs):
        if contains_text in p.text:
            return i
    return None

# ============================================================================
# ТАБЛИЦА 2: Компоненты Unity (после 1.3)
# ============================================================================
table2_data = [
    ["Компонент", "Назначение", "Использование в проекте"],
    ["Transform", "Позиция, вращение, масштаб", "Все игровые объекты"],
    ["SpriteRenderer", "Отрисовка 2D-спрайта", "Персонаж, деревья, камни, NPC"],
    ["Rigidbody2D", "Физическое моделирование", "Персонаж, подбираемые предметы"],
    ["BoxCollider2D", "Прямоугольная коллизия", "Стены, двери, триггеры зон"],
    ["CircleCollider2D", "Круглая коллизия", "Зона подбора лута"],
    ["PolygonCollider2D", "Сложная форма коллизии", "Ландшафт, неровные объекты"],
    ["Animator", "Управление анимациями", "Персонаж, животные, сундуки, время"],
    ["NavMeshAgent2D", "Поиск пути в 2D", "Животные (курицы, коровы)"],
    ["SortingGroup", "Группировка слоёв отрисовки", "Многослойные объекты (дома)"],
    ["Canvas", "Контейнер UI-элементов", "HUD, главное меню, инвентарь"],
]

idx_2 = find_paragraph_idx("Это обеспечивает высокую степень повторного использования кода и упрощает отладку")
if idx_2 is not None:
    doc.paragraphs[idx_2].insert_paragraph_before(
        text="Основные компоненты Unity, используемые в проекте, представлены в таблице 2.",
        style="Normal"
    )
    table2 = doc.add_table(rows=len(table2_data), cols=len(table2_data[0]))
    for i, row_data in enumerate(table2_data):
        for j, cell_text in enumerate(row_data):
            table2.cell(i, j).text = cell_text
    style_table(table2)
    doc.add_paragraph('Таблица 2 — Основные компоненты Unity, используемые в проекте', style='Caption')
    print("✅ Таблица 2 добавлена: Компоненты Unity")
else:
    print("⚠️ Не найден параграф для Таблицы 2")

# ============================================================================
# ТАБЛИЦА 3: ScriptableObject предметы (после 1.4)
# ============================================================================
table3_data = [
    ["Предмет", "Тип", "Действие", "Стакаемость", "Назначение"],
    ["Wood (Дерево)", "BuildingBlock", "Нет", "Да (до 64)", "Ресурс для крафта досок"],
    ["Stone (Камень)", "BuildingBlock", "Нет", "Да (до 64)", "Ресурс для крафта кирки"],
    ["Docka (Доски)", "BuildingBlock", "Нет", "Да (до 64)", "Продукт крафта (1 дерево = 4 доски)"],
    ["Axe (Топор)", "Tool", "Mine", "Нет", "Рубка деревьев"],
    ["PickAxe (Кирка)", "Tool", "Mine", "Нет", "Добыча камня/руды"],
    ["Aplle (Яблоко)", "Food", "Нет", "Да (до 64)", "Восстановление здоровья"],
    ["Kabachok (Кабачок)", "Food", "Нет", "Да (до 64)", "Продукт фермерства"],
    ["Strawberry (Клубника)", "Food", "Нет", "Да (до 64)", "Сбор с кустов"],
]

idx_3 = find_paragraph_idx("без вмешательства программиста")
if idx_3 is not None:
    doc.paragraphs[idx_3].insert_paragraph_before(
        text="Предметы, реализованные через ScriptableObject в проекте, представлены в таблице 3.",
        style="Normal"
    )
    table3 = doc.add_table(rows=len(table3_data), cols=len(table3_data[0]))
    for i, row_data in enumerate(table3_data):
        for j, cell_text in enumerate(row_data):
            table3.cell(i, j).text = cell_text
    style_table(table3)
    doc.add_paragraph('Таблица 3 — Предметы системы инвентаря (ScriptableObject)', style='Caption')
    print("✅ Таблица 3 добавлена: ScriptableObject предметы")
else:
    print("⚠️ Не найден параграф для Таблицы 3")

# ============================================================================
# ТАБЛИЦА 4: Рецепты крафта (после Таблицы 3)
# ============================================================================
table4_data = [
    ["Рецепт", "Сетка ингредиентов (3x3)", "Результат", "Количество"],
    ["Recipe_Boards", "Пусто / Wood / Пусто\nПусто / Пусто / Пусто\nПусто / Пусто / Пусто", "Docka (Доски)", "4"],
    ["Recipe_WoodPickAxe", "Stone / Пусто / Stone\nПусто / Wood / Пусто\nПусто / Wood / Пусто", "PickAxe (Кирка)", "1"],
    ["Recipe_StoneAxe", "Stone / Stone / Stone\nПусто / Wood / Пусто\nПусто / Wood / Пусто", "Axe (Топор)", "1"],
]

idx_4 = find_paragraph_idx("Таблица 3 — Предметы системы инвентаря")
if idx_4 is not None:
    doc.paragraphs[idx_4].insert_paragraph_before(
        text="Рецепты крафта, доступные в игре, представлены в таблице 4.",
        style="Normal"
    )
    table4 = doc.add_table(rows=len(table4_data), cols=len(table4_data[0]))
    for i, row_data in enumerate(table4_data):
        for j, cell_text in enumerate(row_data):
            table4.cell(i, j).text = cell_text
    style_table(table4)
    doc.add_paragraph('Таблица 4 — Рецепты системы крафта', style='Caption')
    print("✅ Таблица 4 добавлена: Рецепты крафта")
else:
    print("⚠️ Не найден параграф для Таблицы 4")

# ============================================================================
# Сохранение
# ============================================================================
output_path = r'C:\UnityGames\GameYear\Diplom\Диплома.docx'
doc.save(output_path)
print(f"\n✅ Файл сохранён: {output_path}")
print(f"Добавлено таблиц: {sum([idx_2 is not None, idx_3 is not None, idx_4 is not None])} из 3")
