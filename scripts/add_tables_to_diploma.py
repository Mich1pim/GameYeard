from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document(r'C:\UnityGames\GameYear\Diplom\Диплома.docx')

def style_table(table):
    """Применяет стили к таблице"""
    # Заголовочная строка
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Закраска заголовка
        shading = cell._tc.get_or_add_tcPr()
        shading_elem = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '2F5496',
            qn('w:val'): 'clear'
        })
        shading.append(shading_elem)
    
    # Остальные строки
    for i, row in enumerate(table.rows[1:], 1):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(10)
            # Чередование фона
            if i % 2 == 0:
                shading = cell._tc.get_or_add_tcPr()
                shading_elem = shading.makeelement(qn('w:shd'), {
                    qn('w:fill'): 'D6E4F0',
                    qn('w:val'): 'clear'
                })
                shading.append(shading_elem)

def add_caption_after_table(doc, table, caption_text):
    """Добавляет подпись после таблицы"""
    last_cell = table.rows[-1].cells[-1]
    p = last_cell._tc.getparent().getparent().getparent()  # получаем paragraph после table
    # Альтернативный подход - вставляем параграф после table
    tbl = table._tbl
    p_after = tbl.getnext()
    if p_after is not None:
        new_p = p_after.makeelement(qn('w:p'), {})
        new_r = new_p.makeelement(qn('w:r'), {})
        new_t = new_r.makeelement(qn('w:t'), {})
        new_t.text = caption_text
        new_r.append(new_t)
        new_p.append(new_r)
        # Вставляем перед p_after
        p_after.getparent().insert(list(p_after.getparent()).index(p_after), new_p)

# ============================================================================
# ТАБЛИЦА 1: Сравнительная характеристика игровых жанров (после 1.1)
# ============================================================================
insert_after_text_1 = "Экономическая система управления хозяйством."
table1_data = [
    ["Жанр", "Ключевые механики", "Примеры", "Применимость к проекту"],
    ["Экшн", "Быстрая реакция, координация", "Платформеры, шутеры", "Атака инструментами"],
    ["Приключения", "Исследование, головоломки", "Zelda, Myst", "Система исследования мира"],
    ["RPG", "Прокачка, сюжет", "Stardew Valley, Terraria", "Экономика, крафт, NPC"],
    ["Стратегии", "Планирование, ресурсы", "Civilization, Age of Empires", "Управление фермой"],
    ["Симуляторы", "Имитация деятельности", "Farming Simulator", "Фермерское хозяйство"],
    ["Гибрид (проект)", "Крафт + исследование + ферма", "Stardew Valley + Terraria", "Все вышеперечисленное"],
]

# Находим параграф с нужным текстом
found_idx_1 = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == insert_after_text_1:
        found_idx_1 = i
        break

if found_idx_1 is not None:
    # Вставляем параграф-заголовок перед таблицей
    doc.paragraphs[found_idx_1].insert_paragraph_before(
        text="Сравнительная характеристика игровых жанров представлена в таблице 1.",
        style="Normal"
    )
    
    # Вставляем таблицу
    table1 = doc.add_table(rows=len(table1_data), cols=len(table1_data[0]))
    for i, row_data in enumerate(table1_data):
        for j, cell_text in enumerate(row_data):
            table1.cell(i, j).text = cell_text
    style_table(table1)
    
    # Подпись таблицы
    doc.add_paragraph('Таблица 1 — Сравнительная характеристика игровых жанров', style='Caption')
    print("✅ Таблица 1 добавлена: Сравнительная характеристика игровых жанров")
else:
    print("⚠️ Не найден параграф для Таблицы 1")

# ============================================================================
# ТАБЛИЦА 2: Компоненты Unity (после раздела 1.3)
# ============================================================================
insert_after_text_2 = "Это обеспечивает высокую степень повторного использования кода и упрощает отладку."
table2_data = [
    ["Компонент", "Назначение", "Использование в проекте"],
    ["Transform", "Позиция, вращение, масштаб", "Все игровые объекты"],
    ["SpriteRenderer", "Отрисовка 2D-спрайта", "Персонаж, деревья, камни, NPC"],
    ["Rigidbody2D", "Физическое моделирование", "Персонаж, подбираемые предметы"],
    ["BoxCollider2D", "Прямоугольная коллизия", "Стены, двери, триггеры зон"],
    ["CircleCollider2D", "Круглая коллизия", "Зона подбора лута"],
    ["PolygonCollider2D", "Сложная форма коллизии", "Ландшафт, неровные объекты"],
    ["Animator", "Управление анимациями", "Персонаж, животные, сундуки, время"],
    ["NavMeshAgent2D", "Поиск пути в 2D", "Животные (курицы, коровы)"],
    ["SortingGroup", "Группировка слоёв отрисовки", "Многослойные объекты (дома)"],
    ["Canvas", "Контейнер UI-элементов", "HUD, главное меню, инвентарь"],
]

found_idx_2 = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == insert_after_text_2:
        found_idx_2 = i
        break

if found_idx_2 is not None:
    doc.paragraphs[found_idx_2].insert_paragraph_before(
        text="Основные компоненты Unity, используемые в проекте, представлены в таблице 2.",
        style="Normal"
    )
    
    table2 = doc.add_table(rows=len(table2_data), cols=len(table2_data[0]))
    for i, row_data in enumerate(table2_data):
        for j, cell_text in enumerate(row_data):
            table2.cell(i, j).text = cell_text
    style_table(table2)
    
    doc.add_paragraph('Таблица 2 — Основные компоненты Unity, используемые в проекте', style='Caption')
    print("✅ Таблица 2 добавлена: Компоненты Unity")
else:
    print("⚠️ Не найден параграф для Таблицы 2")

# ============================================================================
# ТАБЛИЦА 3: ScriptableObject предметы (после раздела 1.4)
# ============================================================================
insert_after_text_3 = "Это позволяет геймдизайнерам балансировать игровые параметры (количество ресурсов в рецептах, характеристики предметов) без вмешательства программиста."
table3_data = [
    ["Предмет", "Тип", "Действие", "Стакаемость", "Назначение"],
    ["Wood (Дерево)", "BuildingBlock", "Нет", "Да (до 64)", "Ресурс для крафта досок"],
    ["Stone (Камень)", "BuildingBlock", "Нет", "Да (до 64)", "Ресурс для крафта кирки"],
    ["Docka (Доски)", "BuildingBlock", "Нет", "Да (до 64)", "Продукт крафта (1 дерево = 4 доски)"],
    ["Axe (Топор)", "Tool", "Mine", "Нет", "Рубка деревьев"],
    ["PickAxe (Кирка)", "Tool", "Mine", "Нет", "Добыча камня/руды"],
    ["Aplle (Яблоко)", "Food", "Нет", "Да (до 64)", "Восстановление здоровья"],
    ["Kabachok (Кабачок)", "Food", "Нет", "Да (до 64)", "Продукт фермерства"],
    ["Strawberry (Клубника)", "Food", "Нет", "Да (до 64)", "Сбор с кустов"],
]

found_idx_3 = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == insert_after_text_3:
        found_idx_3 = i
        break

if found_idx_3 is not None:
    doc.paragraphs[found_idx_3].insert_paragraph_before(
        text="Предметы, реализованные через ScriptableObject в проекте, представлены в таблице 3.",
        style="Normal"
    )
    
    table3 = doc.add_table(rows=len(table3_data), cols=len(table3_data[0]))
    for i, row_data in enumerate(table3_data):
        for j, cell_text in enumerate(row_data):
            table3.cell(i, j).text = cell_text
    style_table(table3)
    
    doc.add_paragraph('Таблица 3 — Предметы системы инвентаря (ScriptableObject)', style='Caption')
    print("✅ Таблица 3 добавлена: ScriptableObject предметы")
else:
    print("⚠️ Не найден параграф для Таблицы 3")

# ============================================================================
# ТАБЛИЦА 4: Рецепты крафта (после Таблицы 3)
# ============================================================================
table4_data = [
    ["Рецепт", "Сетка ингредиентов (3×3)", "Результат", "Количество"],
    ["Recipe_Boards", "Пусто / Wood / Пусто\nПусто / Пусто / Пусто\nПусто / Пусто / Пусто", "Docka (Доски)", "4"],
    ["Recipe_WoodPickAxe", "Stone / Пусто / Stone\nПусто / Wood / Пусто\nПусто / Wood / Пусто", "PickAxe (Кирка)", "1"],
    ["Recipe_StoneAxe", "Stone / Stone / Stone\nПусто / Wood / Пусто\nПусто / Wood / Пусто", "Axe (Топор)", "1"],
]

# Вставляем после последней добавленной таблицы (после Caption Таблицы 3)
found_idx_4 = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "Таблица 3 — Предметы системы инвентаря (ScriptableObject)":
        found_idx_4 = i
        break

if found_idx_4 is not None:
    doc.paragraphs[found_idx_4].insert_paragraph_before(
        text="Рецепты крафта, доступные в игре, представлены в таблице 4.",
        style="Normal"
    )
    
    table4 = doc.add_table(rows=len(table4_data), cols=len(table4_data[0]))
    for i, row_data in enumerate(table4_data):
        for j, cell_text in enumerate(row_data):
            table4.cell(i, j).text = cell_text
    style_table(table4)
    
    doc.add_paragraph('Таблица 4 — Рецепты системы крафта', style='Caption')
    print("✅ Таблица 4 добавлена: Рецепты крафта")
else:
    print("⚠️ Не найден параграф для Таблицы 4")

# ============================================================================
# ТАБЛИЦА 5: Физика 2D vs 3D (после раздела 1.7)
# ============================================================================
insert_after_text_5 = "В проекте система физики 2D обеспечивает перемещение персонажа (Rigidbody2D с типом Dynamic), обнаружение столкновений с препятствиями (BoxCollider2D), подбор предметов через триггеры (Loot.cs, Looting.cs) и взаимодействие с интерактивными объектами (Chest.cs, Door.cs)."
table5_data = [
    ["Параметр", "Physics 2D", "Physics 3D"],
    ["Физический движок", "Box2D", "NVIDIA PhysX"],
    ["Компонент физики", "Rigidbody2D", "Rigidbody"],
    ["Коллайдеры", "BoxCollider2D, CircleCollider2D, PolygonCollider2D", "BoxCollider, SphereCollider, MeshCollider"],
    ["Оси координат", "X, Y (плоскость)", "X, Y, Z (объём)"],
    ["События столкновений", "OnCollisionEnter2D, OnTriggerEnter2D", "OnCollisionEnter, OnTriggerEnter"],
    ["Производительность", "Выше (меньше вычислений)", "Ниже (больше вычислений)"],
    ["Применение", "Пиксельные, изометрические, платформеры", "3D-шутеры, гонки, симуляторы"],
    ["В проекте", "Используется ✅", "Не используется ❌"],
]

found_idx_5 = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == insert_after_text_5:
        found_idx_5 = i
        break

if found_idx_5 is not None:
    doc.paragraphs[found_idx_5].insert_paragraph_before(
        text="Сравнение систем физики представлено в таблице 5.",
        style="Normal"
    )
    
    table5 = doc.add_table(rows=len(table5_data), cols=len(table5_data[0]))
    for i, row_data in enumerate(table5_data):
        for j, cell_text in enumerate(row_data):
            table5.cell(i, j).text = cell_text
    style_table(table5)
    
    doc.add_paragraph('Таблица 5 — Сравнение систем физики 2D и 3D в Unity', style='Caption')
    print("✅ Таблица 5 добавлена: Физика 2D vs 3D")
else:
    print("⚠️ Не найден параграф для Таблицы 5")

# ============================================================================
# ТАБЛИЦА 6: Системные требования (в раздел 2.1.5)
# ============================================================================
insert_after_text_6 = "Место на диске: 2 GB."
table6_data = [
    ["Параметр", "Минимальные", "Рекомендуемые"],
    ["Операционная система", "Windows 10 x64", "Windows 11 x64"],
    ["Процессор", "Intel Core i3 / AMD Ryzen 3", "Intel Core i5 / AMD Ryzen 5"],
    ["Оперативная память", "4 GB RAM", "8 GB RAM"],
    ["Видеокарта", "DirectX 11 compatible", "DirectX 12, 2 GB VRAM"],
    ["Место на диске", "2 GB", "4 GB (с запасом для сохранений)"],
    ["Разрешение экрана", "1280×720 (HD)", "1920×1080 (Full HD)"],
]

found_idx_6 = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == insert_after_text_6:
        found_idx_6 = i
        break

if found_idx_6 is not None:
    doc.paragraphs[found_idx_6].insert_paragraph_before(
        text="Системные требования представлены в таблице 6.",
        style="Normal"
    )
    
    table6 = doc.add_table(rows=len(table6_data), cols=len(table6_data[0]))
    for i, row_data in enumerate(table6_data):
        for j, cell_text in enumerate(row_data):
            table6.cell(i, j).text = cell_text
    style_table(table6)
    
    doc.add_paragraph('Таблица 6 — Системные требования программного продукта', style='Caption')
    print("✅ Таблица 6 добавлена: Системные требования")
else:
    print("⚠️ Не найден параграф для Таблицы 6")

# ============================================================================
# Сохранение
# ============================================================================
output_path = r'C:\UnityGames\GameYear\Diplom\Диплома.docx'
doc.save(output_path)
print(f"\n✅ Файл сохранён: {output_path}")
print(f"Всего добавлено таблиц: 6")
